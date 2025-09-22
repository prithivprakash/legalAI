"""
Backend logic for the Legal Document Simplifier AI.

Implements:
- parse_file(file) -> extract text for pdf and docx
- chunk_and_store(...) -> chunk using RecursiveCharacterTextSplitter, embed (HuggingFace), store in Chroma
- semantic_search(...) -> embed query, retrieve top_k chunks from Chroma, call Google Gemini with retrieved context + query
- summarize_text(text) -> use Google Gemini to produce plain-language summary
- analyze_document_for_risks(text) -> returns (risks_dict, obligations_list)
- compare_documents(text1, text2) -> simple redline/diff output as markdown

Requirements:
pip install langchain chromadb sentence-transformers transformers pdfplumber python-docx numpy
pip install langchain-google-genai
"""
import json
import os
import tempfile
import uuid
import difflib
from typing import Tuple, Dict, List, Any
from datetime import datetime, timedelta
import re
from dateutil.parser import parse as parse_date
from dateutil.relativedelta import relativedelta
from datetime import datetime
import re
from dateutil.parser import parse as parse_date


# Environment / persistence
from dotenv import load_dotenv

load_dotenv()  # load .env in working directory

VECTORSTORE_PERSIST_DIR = os.getenv("VECTORSTORE_PERSIST_DIR", "./chroma_db")

# LangChain / Embeddings / Vector DB / LLM
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import Chroma
from langchain.docstore.document import Document
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain

# File parsing
import pdfplumber
import docx
import io
import numpy as np
def simple_date_validator(text: str) -> dict:
    """
    Simple date extraction and validation - safe to add without conflicts
    Returns current validity status of documents
    """
    current_date = datetime.now()
    
    # Common date patterns to look for
    date_patterns = [
        r'(\d{1,2}[-/]\d{1,2}[-/]\d{4})',  # dd/mm/yyyy or mm/dd/yyyy
        r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})',  # yyyy/mm/dd
        r'(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})',  # dd Mon yyyy
        r'((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})'  # Month dd, yyyy
    ]
    
    results = {
        "status": "UNKNOWN",
        "message": "Unable to determine validity",
        "dates_found": [],
        "is_valid": None,
        "action_needed": []
    }
    
    text_lower = text.lower()
    
    # Look for dates with context
    all_found_dates = []
    
    for pattern in date_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            date_str = match.group(1)
            # Get context around the date
            start = max(0, match.start() - 40)
            end = min(len(text), match.end() + 40)
            context = text[start:end].lower()
            
            try:
                # First try with dateutil parser (more flexible)
                try:
                    parsed_date = parse_date(date_str, fuzzy=True)
                except:
                    # Fallback to manual parsing with more format attempts
                    date_str_clean = date_str.replace('-', '/').replace('.', '/')
                    formats = ['%d/%m/%Y', '%m/%d/%Y', '%Y/%m/%d', '%Y-%m-%d', '%d-%m-%Y', '%m-%d-%Y']
                    
                    for fmt in formats:
                        try:
                            test_fmt = fmt.replace('-', '/').replace('.', '/')
                            parsed_date = datetime.strptime(date_str_clean, test_fmt)
                            break
                        except:
                            continue
                    else:
                        # Last resort: manual parsing
                        parts = re.split(r'[-/.\s]', date_str)
                        if len(parts) >= 3:
                            try:
                                # Try DD/MM/YYYY first, then MM/DD/YYYY, then YYYY/MM/DD
                                day, month, year = int(parts[0]), int(parts[1]), int(parts[2])
                                if year < 100:
                                    year += 2000
                                if year < 1900 or year > 2100:
                                    continue
                                parsed_date = datetime(year, month, day)
                            except:
                                try:
                                    month, day, year = int(parts[0]), int(parts[1]), int(parts[2])
                                    if year < 100:
                                        year += 2000
                                    parsed_date = datetime(year, month, day)
                                except:
                                    continue
                
                # Calculate days difference
                days_diff = (parsed_date - current_date).days
                
                # Determine if this is an expiry date or issue date
                is_expiry = any(word in context for word in ['expir', 'valid until', 'valid till', 'expire'])
                is_issue = any(word in context for word in ['issue', 'from', 'effective', 'start'])
                
                all_found_dates.append({
                    'date': parsed_date.strftime('%Y-%m-%d'),
                    'original': date_str,
                    'days_from_now': days_diff,
                    'is_expiry': is_expiry,
                    'is_issue': is_issue,
                    'context': context
                })
                
            except:
                continue
    
    # Analyze findings
    if all_found_dates:
        results["dates_found"] = all_found_dates
        
        # Look for expiry dates first
        expiry_dates = [d for d in all_found_dates if d['is_expiry']]
        
        if expiry_dates:
            # Use the earliest expiry date
            earliest_expiry = min(expiry_dates, key=lambda x: x['days_from_now'])
            days_remaining = earliest_expiry['days_from_now']
            
            if days_remaining < 0:
                results["status"] = "EXPIRED"
                results["is_valid"] = False
                results["message"] = f"Document EXPIRED {abs(days_remaining)} days ago on {earliest_expiry['date']}"
                results["action_needed"] = ["Immediate renewal required", "Document is currently NOT VALID"]
            elif days_remaining <= 30:
                results["status"] = "EXPIRING_SOON" 
                results["is_valid"] = True
                results["message"] = f"Document expires in {days_remaining} days on {earliest_expiry['date']}"
                results["action_needed"] = ["Start renewal process immediately", f"Only {days_remaining} days remaining"]
            else:
                results["status"] = "VALID"
                results["is_valid"] = True
                results["message"] = f"Document is valid until {earliest_expiry['date']} ({days_remaining} days remaining)"
                results["action_needed"] = []
        
        else:
            # No explicit expiry dates found, try to estimate from issue dates
            issue_dates = [d for d in all_found_dates if d['is_issue']]
            
            if issue_dates:
                latest_issue = max(issue_dates, key=lambda x: x['days_from_now'])
                issue_date = datetime.strptime(latest_issue['date'], '%Y-%m-%d')
                
                # Estimate based on document type
                validity_months = 12  # Default to 1 year
                
                if any(word in text_lower for word in ['puc', 'pollution']):
                    validity_months = 6  # PUC is 6 months
                elif any(word in text_lower for word in ['fitness']):
                    validity_months = 12  # Fitness is 1 year
                elif any(word in text_lower for word in ['license', 'licence']):
                    validity_months = 12  # Most licenses are 1 year
                
                estimated_expiry = issue_date + timedelta(days=validity_months*30)
                days_until_estimated_expiry = (estimated_expiry - current_date).days
                
                if days_until_estimated_expiry < 0:
                    results["status"] = "LIKELY_EXPIRED"
                    results["is_valid"] = False
                    results["message"] = f"Likely EXPIRED (estimated expiry: {estimated_expiry.strftime('%Y-%m-%d')})"
                    results["action_needed"] = ["Verify current status with issuer", "Likely renewal required"]
                else:
                    results["status"] = "LIKELY_VALID"
                    results["is_valid"] = True
                    results["message"] = f"Likely valid until approximately {estimated_expiry.strftime('%Y-%m-%d')}"
                    results["action_needed"] = ["Verify exact expiry date with issuer"]
    
    return results

# Enhanced semantic search that uses the date validator
def semantic_search_with_dates(
    query: str,
    collection_name: str = "legal_docs",
    top_k: int = 5,
    embedding_model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
    persist_directory: str = VECTORSTORE_PERSIST_DIR,
    model_name: str = None,
) -> str:
    """
    Enhanced semantic search that includes date validation for better answers
    """
    # Check if query is asking about validity/expiry
    query_lower = query.lower()
    is_validity_query = any(word in query_lower for word in [
        'valid', 'expired', 'expiry', 'current', 'active', 'still good'
    ])
    
    # Get standard semantic search result first
    try:
        from backend import semantic_search
        base_answer = semantic_search(query, collection_name, top_k, embedding_model_name, persist_directory, model_name)
    except:
        # If semantic_search doesn't exist, do basic retrieval
        embedder = _get_embedding_model(embedding_model_name)
        vectordb = Chroma(persist_directory=persist_directory, embedding_function=embedder, collection_name=collection_name)
        docs = vectordb.similarity_search(query, k=top_k)
        full_text = "\n".join([doc.page_content for doc in docs])
        base_answer = f"Based on document content: {full_text[:1000]}..."
    
    # If it's a validity query, enhance with date validation
    if is_validity_query:
        try:
            # Get the full document text for date analysis
            embedder = _get_embedding_model(embedding_model_name)
            vectordb = Chroma(persist_directory=persist_directory, embedding_function=embedder, collection_name=collection_name)
            docs = vectordb.similarity_search(query, k=top_k)
            full_text = "\n".join([doc.page_content for doc in docs])
            
            # Run date validation
            date_results = simple_date_validator(full_text)
            
            # Generate enhanced answer based on date validation
            if date_results["status"] != "UNKNOWN":
                current_date = datetime.now().strftime("%B %d, %Y")
                
                if date_results["status"] == "EXPIRED":
                    enhanced_answer = f"""**🔴 NO - Your document is NOT VALID**

**Current Status:** EXPIRED
**Date Analysis:** {date_results['message']}
**Validity:** ❌ NOT VALID (as of {current_date})

**Immediate Actions Required:**
"""
                    for action in date_results["action_needed"]:
                        enhanced_answer += f"• {action}\n"
                    
                    enhanced_answer += f"\n**Original Analysis:** {base_answer}"
                    
                elif date_results["status"] == "VALID":
                    enhanced_answer = f"""**🟢 YES - Your document is VALID**

**Current Status:** VALID
**Date Analysis:** {date_results['message']}
**Validity:** ✅ CURRENTLY VALID (as of {current_date})

**No immediate action required** - document is currently valid.

**Original Analysis:** {base_answer}"""

                elif date_results["status"] == "EXPIRING_SOON":
                    enhanced_answer = f"""**🟡 YES - Your document is VALID but expiring soon**

**Current Status:** EXPIRING SOON
**Date Analysis:** {date_results['message']}
**Validity:** ⚠️ VALID but action needed (as of {current_date})

**Actions Needed:**
"""
                    for action in date_results["action_needed"]:
                        enhanced_answer += f"• {action}\n"
                    
                    enhanced_answer += f"\n**Original Analysis:** {base_answer}"
                
                elif "LIKELY" in date_results["status"]:
                    is_likely_valid = "VALID" in date_results["status"]
                    status_icon = "🟢" if is_likely_valid else "🔴"
                    status_text = "LIKELY VALID" if is_likely_valid else "LIKELY EXPIRED"
                    
                    enhanced_answer = f"""**{status_icon} {status_text} - Verification Recommended**

**Estimated Status:** {status_text}
**Date Analysis:** {date_results['message']}
**Validity:** {"⚠️ Likely valid but verify" if is_likely_valid else "❌ Likely expired"}

**Recommended Actions:**
"""
                    for action in date_results["action_needed"]:
                        enhanced_answer += f"• {action}\n"
                    
                    enhanced_answer += f"\n**Original Analysis:** {base_answer}"
                
                return enhanced_answer
            
        except Exception as e:
            print(f"Date validation failed: {e}")
    
    # Return original answer if not validity query or if date validation failed
    return base_answer
# Configure Gemini via environment variables
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-1.5-flash")  # default model

# Basic guard
if not GEMINI_API_KEY:
    print("Warning: Gemini API key not set. LLM calls will fail without it.")

def parse_file(uploaded_file) -> str:
    """
    Parse uploaded pdf/docx file-like object and return extracted text.
    uploaded_file is a Streamlit UploadedFile (file-like).
    """
    fname = uploaded_file.name.lower()
    # Move to bytesIO to ensure compatibility
    content = uploaded_file.read()
    text = ""
    if fname.endswith(".pdf"):
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages = []
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
            text = "\n\n".join(pages)
    elif fname.endswith(".docx"):
        doc = docx.Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs]
        text = "\n\n".join(paragraphs)
    else:
        # fallback
        text = content.decode("utf-8", errors="ignore")
    return text

def chunk_and_store(
    text: str,
    collection_name: str = "legal_docs",
    chunk_size: int = 500,
    chunk_overlap: int = 50,
    embedding_model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
    persist_directory: str = VECTORSTORE_PERSIST_DIR,
):
    """
    1) Chunk text using RecursiveCharacterTextSplitter
    2) Create embeddings via HuggingFaceEmbeddings
    3) Store chunks into Chroma (persisted to persist_directory)
    """
    # 1. Chunk
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " ", ""],
    )
    chunks = splitter.split_text(text)
    docs = []
    for i, c in enumerate(chunks):
        md = {"chunk_id": f"{i}", "source": f"{collection_name}"}
        docs.append(Document(page_content=c, metadata=md))

    # 2. Embeddings
    embedder = _get_embedding_model(embedding_model_name)

    # 3. Chroma store (langchain wrapper)
    chroma_db = Chroma.from_documents(
        documents=docs,
        embedding=embedder,
        persist_directory=persist_directory,
        collection_name=collection_name,
    )
    chroma_db.persist()
    return chroma_db

def semantic_search(
    query: str,
    collection_name: str = "legal_docs",
    top_k: int = 5,
    embedding_model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
    persist_directory: str = VECTORSTORE_PERSIST_DIR,
    model_name: str = None,
) -> str:
    """
    High-level RAG:
      - embed query (HuggingFace)
      - retrieve top_k similar chunks from Chroma
      - compose a context prompt including retrieved chunks and the user query
      - call Google Gemini to answer concisely in plain language

    Returns final LLM string answer.
    """
    # 1) Load vector store
    embedder = _get_embedding_model(embedding_model_name)
    vectordb = Chroma(persist_directory=persist_directory, embedding_function=embedder, collection_name=collection_name)

    # 2) Retrieval
    docs = vectordb.similarity_search(query, k=top_k)

    # join retrieved docs into a context string (limit size if necessary)
    ctxs = []
    for i, d in enumerate(docs):
        # include small metadata reference
        source = d.metadata.get("source", "")
        ctxs.append(f"--- chunk {i} from {source} ---\n{d.page_content.strip()[:2000]}")  # limit 2000 chars per chunk

    context_text = "\n\n".join(ctxs)
    # 3) LLM call
    llm = _get_gemini_llm(model_name=model_name, temperature=0.0)

    prompt_template = """You are a legal-document assistant. Use the following retrieved document chunks (below) to answer the user's question.
Do NOT hallucinate facts — if the answer is not present in the retrieved chunks, say you don't know or ask for clarification.

RETRIEVED CHUNKS:
{context}

USER QUESTION:
{question}

INSTRUCTIONS:
- Answer in plain, simple English suitable for a non-lawyer.
- When listing obligations or risks, label them clearly and make them concise.
- If you refer to a clause, quote the relevant short phrase (<= 50 words) and indicate which chunk it came from.
- Provide actionable next steps the user could take (2-4 bullets).

Provide the final answer below:
"""

    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = LLMChain(llm=llm, prompt=prompt)
    resp = chain.run({"context": context_text, "question": query})
    return resp


def summarize_text(text: str, model_name: str = None) -> str:
    """
    Simple plain-language summary using Google Gemini model.
    """
    llm = _get_gemini_llm(model_name=model_name, temperature=0.0)
    prompt_template = """You are an expert legal simplifier. Summarize the document in plain language for a non-lawyer.
Keep it under ~250 words. Highlight: (1) Purpose of the document, (2) Key obligations, (3) Important dates/terms, (4) Any immediate actions recommended.
DOCUMENT:
{doc}
"""
    prompt = PromptTemplate(template=prompt_template, input_variables=["doc"])
    chain = LLMChain(llm=llm, prompt=prompt)
    return chain.run({"doc": text})



def analyze_document_for_risks(text: str, model_name: str = None) -> Tuple[Dict[str, List[str]], List[str]]:
    """
    Use the LLM to identify risks and obligations in any legal document. Returns:
      - risks: dict with keys "High", "Medium", "Low" -> list of strings
      - obligations: list of strings
    This is a comprehensive implementation for all legal document types.
    """
    llm = _get_gemini_llm(model_name=model_name, temperature=0.0)
    
    # Simplified prompt that's more likely to produce valid JSON
    prompt_template = """You are a legal analyst. Analyze this document and identify risks and obligations.

DOCUMENT:
{doc}

Respond with ONLY valid JSON in this exact format:
{{
  "High": ["high risk item 1", "high risk item 2"],
  "Medium": ["medium risk item 1", "medium risk item 2"],
  "Low": ["low risk item 1"],
  "Obligations": ["obligation 1", "obligation 2"]
}}

HIGH RISKS: expired items, breach clauses, penalties, non-compliance, overdue payments
MEDIUM RISKS: upcoming deadlines, performance requirements, insurance needs, reporting duties
LOW RISKS: standard clauses, routine provisions
OBLIGATIONS: specific actions required, payments due, compliance requirements

JSON response:"""

    prompt = PromptTemplate(template=prompt_template, input_variables=["doc"])
    
    # Initialize default values
    risks = {"High": [], "Medium": [], "Low": []}
    obligations = []
    
    try:
        # Get LLM response
        chain = LLMChain(llm=llm, prompt=prompt)
        raw_response = chain.run({"doc": text[:8000]})  # Limit text size
        
        print(f"Raw LLM Response: {raw_response}")  # Debug print
        
        # Clean the response
        cleaned_response = raw_response.strip()
        
        # Remove markdown code blocks if present
        if cleaned_response.startswith('```'):
            lines = cleaned_response.split('\n')
            # Remove first and last lines if they're markdown markers
            if lines[0].startswith('```'):
                lines = lines[1:]
            if lines and lines[-1].strip() == '```':
                lines = lines[:-1]
            cleaned_response = '\n'.join(lines)
        
        # Remove any text before the first { and after the last }
        import re
        json_match = re.search(r'\{.*\}', cleaned_response, re.DOTALL)
        if json_match:
            cleaned_response = json_match.group(0)
        
        print(f"Cleaned Response: {cleaned_response}")  # Debug print
        
        # Try to parse JSON
        
        parsed_result = json.loads(cleaned_response)
        
        # Safely extract data
        for risk_level in ["High", "Medium", "Low"]:
            if risk_level in parsed_result:
                if isinstance(parsed_result[risk_level], list):
                    risks[risk_level] = [str(item).strip() for item in parsed_result[risk_level] if str(item).strip()]
                elif isinstance(parsed_result[risk_level], str) and parsed_result[risk_level].strip():
                    risks[risk_level] = [parsed_result[risk_level].strip()]
        
        if "Obligations" in parsed_result:
            if isinstance(parsed_result["Obligations"], list):
                obligations = [str(item).strip() for item in parsed_result["Obligations"] if str(item).strip()]
            elif isinstance(parsed_result["Obligations"], str) and parsed_result["Obligations"].strip():
                obligations = [parsed_result["Obligations"].strip()]
                
    except Exception as e:
        print(f"JSON parsing failed: {e}")
        
        # Fallback: Use a simpler, more reliable approach
        try:
            # Try alternative prompt that's more structured
            fallback_prompt = f"""Analyze this legal document for risks and obligations. 
            
Document: {text[:5000]}

List HIGH RISKS (critical issues):
List MEDIUM RISKS (important issues):  
List LOW RISKS (minor issues):
List OBLIGATIONS (required actions):

Be specific and concise."""
            
            fallback_chain = LLMChain(llm=llm, prompt=PromptTemplate(
                template=fallback_prompt, input_variables=[]
            ))
            fallback_response = fallback_chain.run({})
            
            print(f"Fallback Response: {fallback_response}")
            
            # Parse the structured response
            lines = fallback_response.split('\n')
            current_section = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                line_lower = line.lower()
                
                # Detect section headers
                if 'high risk' in line_lower:
                    current_section = "High"
                    continue
                elif 'medium risk' in line_lower:
                    current_section = "Medium"
                    continue
                elif 'low risk' in line_lower:
                    current_section = "Low"
                    continue
                elif 'obligation' in line_lower:
                    current_section = "Obligations"
                    continue
                
                # Add content to current section
                if current_section and line:
                    # Clean up bullet points and numbering
                    clean_line = re.sub(r'^[-•*\d\.\s]+', '', line).strip()
                    if clean_line and len(clean_line) > 3:
                        if current_section in risks:
                            risks[current_section].append(clean_line)
                        elif current_section == "Obligations":
                            obligations.append(clean_line)
                            
        except Exception as fallback_error:
            print(f"Fallback parsing also failed: {fallback_error}")
            
            # Last resort: keyword-based analysis
            text_lower = text.lower()
            
            # High risk keywords
            if any(keyword in text_lower for keyword in ['expir', 'breach', 'default', 'penalty', 'terminate']):
                if 'expir' in text_lower:
                    risks["High"].append("Document contains expired or expiring items")
                if 'breach' in text_lower or 'default' in text_lower:
                    risks["High"].append("Breach or default provisions identified")
                if 'penalty' in text_lower:
                    risks["High"].append("Penalty clauses present")
            
            # Medium risk keywords  
            if any(keyword in text_lower for keyword in ['due', 'payment', 'report', 'comply', 'insurance']):
                if 'payment' in text_lower or 'due' in text_lower:
                    risks["Medium"].append("Payment or financial obligations present")
                if 'report' in text_lower:
                    risks["Medium"].append("Reporting requirements identified")
                if 'insurance' in text_lower:
                    risks["Medium"].append("Insurance requirements mentioned")
            
            # Basic obligations
            if any(keyword in text_lower for keyword in ['shall', 'must', 'required', 'obligated']):
                obligations.append("Document contains mandatory obligations requiring review")
    
    # Ensure we always return something useful
    if not any(risks.values()) and not obligations:
        risks["Medium"].append("Document requires professional legal review")
        obligations.append("Review document with legal counsel for specific requirements")
    
    print(f"Final risks: {risks}")  # Debug print
    print(f"Final obligations: {obligations}")  # Debug print
    
    return risks, obligations


# Also add this improved error handling to your frontend section:

def safe_analyze_document_for_risks(text: str, model_name: str = None) -> Tuple[Dict[str, List[str]], List[str]]:
    """
    Wrapper function with comprehensive error handling
    """
    try:
        return analyze_document_for_risks(text, model_name)
    except Exception as e:
        print(f"Analysis failed with error: {e}")
        
        # Return basic analysis based on keywords
        risks = {"High": [], "Medium": [], "Low": []}
        obligations = []
        
        text_lower = text.lower()
        
        # Basic keyword detection
        if 'expir' in text_lower:
            risks["High"].append("Document may contain expired items - manual review required")
        if 'breach' in text_lower or 'default' in text_lower:
            risks["High"].append("Breach or default language detected")
        if 'payment' in text_lower or 'fee' in text_lower:
            risks["Medium"].append("Financial obligations present")
        if 'comply' in text_lower or 'regulation' in text_lower:
            risks["Medium"].append("Compliance requirements indicated")
        
        obligations.append("Manual legal review recommended due to analysis limitations")
        
        return risks, obligations

def safe_analyze_document_for_risks(text: str, model_name: str = None) -> Tuple[Dict[str, List[str]], List[str]]:
    """
    Wrapper function with comprehensive error handling for risk analysis
    """
    try:
        return analyze_document_for_risks(text, model_name)
    except Exception as e:
        print(f"Analysis failed with error: {e}")
        
        # Return basic analysis based on keywords as fallback
        risks = {"High": [], "Medium": [], "Low": []}
        obligations = []
        
        text_lower = text.lower()
        
        # Basic keyword detection for common legal issues
        if 'expir' in text_lower:
            risks["High"].append("Document may contain expired items - manual review required")
        if any(word in text_lower for word in ['breach', 'default', 'violation']):
            risks["High"].append("Breach or default language detected")
        if any(word in text_lower for word in ['overdue', 'past due', 'delinquent']):
            risks["High"].append("Overdue obligations may be present")
        if any(word in text_lower for word in ['penalty', 'fine', 'liquidated damages']):
            risks["High"].append("Penalty clauses identified")
            
        if any(word in text_lower for word in ['payment', 'fee', 'cost', 'deposit']):
            risks["Medium"].append("Financial obligations present")
        if any(word in text_lower for word in ['comply', 'regulation', 'requirement']):
            risks["Medium"].append("Compliance requirements indicated")
        if any(word in text_lower for word in ['report', 'filing', 'submission']):
            risks["Medium"].append("Reporting obligations may apply")
        if any(word in text_lower for word in ['insurance', 'bond', 'guarantee']):
            risks["Medium"].append("Insurance or security requirements mentioned")
            
        # Standard provisions
        if any(word in text_lower for word in ['agreement', 'contract', 'terms']):
            risks["Low"].append("Standard contractual provisions present")
            
        # Basic obligations
        if any(word in text_lower for word in ['shall', 'must', 'required', 'obligated']):
            obligations.append("Document contains mandatory obligations - detailed review recommended")
        
        obligations.append("Analysis encountered technical issues - manual legal review strongly recommended")
        
        return risks, obligations
    
# Date processing
try:
    from dateutil.parser import parse as parse_date
    from dateutil.relativedelta import relativedelta
except ImportError:
    print("Warning: dateutil not installed. Install with: pip install python-dateutil")
    def parse_date(date_str, fuzzy=True):
        # Fallback basic date parsing
        return datetime.strptime(date_str.split()[0], "%Y-%m-%d")

# Configure Gemini
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-1.5-flash")

if not GEMINI_API_KEY:
    print("Warning: Gemini API key not set. LLM calls will fail without it.")

# Time calculation utilities
def calculate_time_difference(start_date: datetime, end_date: datetime = None) -> Dict[str, Any]:
    """
    Calculate comprehensive time difference between dates
    """
    if end_date is None:
        end_date = datetime.now()
    
    diff = end_date - start_date
    total_days = diff.days
    
    # Calculate years, months, days
    years = total_days // 365
    remaining_days = total_days % 365
    months = remaining_days // 30
    days = remaining_days % 30
    
    # Status determination
    if total_days < 0:
        status = "EXPIRED"
        urgency = "CRITICAL"
    elif total_days <= 7:
        status = "EXPIRING_VERY_SOON"
        urgency = "CRITICAL"
    elif total_days <= 30:
        status = "EXPIRING_SOON"
        urgency = "HIGH"
    elif total_days <= 90:
        status = "EXPIRING_MODERATE"
        urgency = "MEDIUM"
    else:
        status = "VALID"
        urgency = "LOW"
    
    return {
        "total_days": total_days,
        "years": years,
        "months": months,
        "days": days,
        "status": status,
        "urgency": urgency,
        "human_readable": f"{years} years, {months} months, {days} days" if years > 0 else f"{months} months, {days} days" if months > 0 else f"{days} days"
    }

def get_typical_validity_periods() -> Dict[str, Dict[str, int]]:
    """
    Return typical validity periods for different document types (in months)
    """
    return {
        "insurance": {
            "motor_insurance": 12,
            "health_insurance": 12,
            "life_insurance": 12,
            "general_insurance": 12,
            "vehicle_insurance": 12
        },
        "licenses": {
            "driving_license": 240,  # 20 years
            "professional_license": 36,  # 3 years
            "business_license": 12,
            "trade_license": 12,
            "liquor_license": 12
        },
        "certificates": {
            "puc_certificate": 6,  # Pollution Under Control
            "fitness_certificate": 12,
            "safety_certificate": 12,
            "quality_certificate": 24,
            "iso_certificate": 36
        },
        "permits": {
            "construction_permit": 24,
            "environmental_permit": 60,
            "import_permit": 6,
            "export_permit": 6
        },
        "registrations": {
            "company_registration": 0,  # Perpetual
            "trademark_registration": 120,  # 10 years
            "patent_registration": 240,  # 20 years
            "vehicle_registration": 180  # 15 years
        }
    }

def intelligent_date_extraction_and_validation(text: str, model_name: str = None) -> Dict[str, Any]:
    """
    Enhanced date extraction with LLM intelligence and time calculation
    """
    llm = _get_gemini_llm(model_name=model_name, temperature=0.0)
    current_date = datetime.now()
    current_date_str = current_date.strftime("%Y-%m-%d")
    
    # Enhanced prompt with parametric knowledge integration
    prompt_template = """You are an expert document analyst with comprehensive knowledge of legal document validity periods.

CURRENT DATE: {current_date}
TODAY IS: {day_of_week}

DOCUMENT TO ANALYZE:
{doc}

TASK: Extract ALL dates and determine document validity status using both the document content AND your knowledge of typical validity periods.

KNOWLEDGE BASE - Typical Validity Periods:
- Motor/Vehicle Insurance: 1 year
- Health Insurance: 1 year  
- PUC Certificate: 6 months
- Fitness Certificate: 1 year
- Driving License: 20 years
- Professional Licenses: 1-3 years
- Business Licenses: 1 year
- ISO Certificates: 3 years
- Construction Permits: 2 years

INSTRUCTIONS:
1. Find ALL dates in the document (issue dates, expiry dates, effective dates)
2. Identify document type (insurance policy, license, certificate, permit, etc.)
3. For each date, calculate exact days remaining until expiry
4. If no expiry date found, estimate based on issue date + typical validity period
5. Determine current validity status: VALID / EXPIRED / EXPIRING_SOON / UNKNOWN
6. Provide specific recommendations based on urgency

Respond with ONLY valid JSON:
{{
  "document_analysis": {{
    "document_type": "specific type identified",
    "confidence_score": "high/medium/low",
    "current_status": "VALID/EXPIRED/EXPIRING_SOON/UNKNOWN",
    "status_reason": "explanation of how status was determined"
  }},
  "dates_extracted": [
    {{
      "date_found": "YYYY-MM-DD",
      "original_text": "exact text from document",
      "date_type": "issue/expiry/effective/renewal",
      "relates_to": "what document/item",
      "days_from_today": number,
      "status": "VALID/EXPIRED/EXPIRING_SOON",
      "urgency_level": "LOW/MEDIUM/HIGH/CRITICAL"
    }}
  ],
  "validity_assessment": {{
    "is_currently_valid": true or false,
    "days_remaining": number or null,
    "expires_on": "YYYY-MM-DD or null",
    "estimated_expiry": "YYYY-MM-DD if estimated from typical periods",
    "confidence_in_assessment": "high/medium/low"
  }},
  "recommendations": [
    "specific actionable recommendation 1",
    "specific actionable recommendation 2"
  ],
  "critical_issues": [
    "list any critical time-sensitive issues"
  ]
}}

JSON response:"""

    prompt = PromptTemplate(template=prompt_template, input_variables=["doc", "current_date", "day_of_week"])
    
    try:
        chain = LLMChain(llm=llm, prompt=prompt)
        raw_response = chain.run({
            "doc": text[:10000], 
            "current_date": current_date_str,
            "day_of_week": current_date.strftime("%A, %B %d, %Y")
        })
        
        print(f"Intelligent Date Analysis Response: {raw_response}")
        
        # Clean and parse JSON
        cleaned_response = raw_response.strip()
        if cleaned_response.startswith('```'):
            lines = cleaned_response.split('\n')
            if lines[0].startswith('```'):
                lines = lines[1:]
            if lines and lines[-1].strip() == '```':
                lines = lines[:-1]
            cleaned_response = '\n'.join(lines)
        
        json_match = re.search(r'\{.*\}', cleaned_response, re.DOTALL)
        if json_match:
            cleaned_response = json_match.group(0)
        
        parsed_result = json.loads(cleaned_response)
        
        # Enhance with calculated time differences
        if "dates_extracted" in parsed_result:
            for date_info in parsed_result["dates_extracted"]:
                try:
                    date_str = date_info.get("date_found")
                    if date_str:
                        parsed_date = datetime.strptime(date_str, "%Y-%m-%d")
                        time_calc = calculate_time_difference(parsed_date, current_date)
                        date_info["time_calculation"] = time_calc
                except Exception:
                    continue
        
        return parsed_result
        
    except Exception as e:
        print(f"Intelligent date analysis failed: {e}")
        return _fallback_intelligent_validation(text, current_date)

def _fallback_intelligent_validation(text: str, current_date: datetime) -> Dict[str, Any]:
    """
    Fallback validation using regex patterns and typical validity periods
    """
    text_lower = text.lower()
    
    # Document type detection
    document_type = "unknown"
    typical_validity_months = 12  # default
    
    # Detect document type and set expected validity
    type_indicators = {
        "insurance": (["insurance", "policy", "coverage", "premium"], 12),
        "puc_certificate": (["puc", "pollution", "emission"], 6),
        "fitness_certificate": (["fitness", "roadworthiness"], 12),
        "driving_license": (["driving", "license", "dl"], 240),
        "vehicle_registration": (["registration", "rc", "vehicle"], 180),
        "professional_license": (["professional", "practice"], 36),
        "business_license": (["business", "trade"], 12)
    }
    
    for doc_type, (keywords, validity_months) in type_indicators.items():
        if any(keyword in text_lower for keyword in keywords):
            document_type = doc_type
            typical_validity_months = validity_months
            break
    
    # Extract dates
    dates_found = []
    date_patterns = [
        r'(\d{1,2}[-/.]\d{1,2}[-/.]\d{4})',  # dd/mm/yyyy, dd-mm-yyyy, dd.mm.yyyy
        r'(\d{4}[-/.]\d{1,2}[-/.]\d{1,2})',  # yyyy/mm/dd, yyyy-mm-dd
        r'(\d{1,2}[-/.]\d{1,2}[-/.]\d{2})',  # dd/mm/yy, dd-mm-yy
        r'(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})',  # dd Mon yyyy
        r'((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})',  # Month dd, yyyy
        r'(\d{2}-\d{2}-\d{4})',  # dd-mm-yyyy
        r'(\d{4}-\d{2}-\d{2})',  # yyyy-mm-dd (ISO format)
        r'(\d{1,2}/\d{1,2}/\d{2})'   # mm/dd/yy or dd/mm/yy
    ]
    
    all_dates = []
    for pattern in date_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            try:
                parsed_date = parse_date(match.group(1), fuzzy=True)
                context_start = max(0, match.start() - 30)
                context_end = min(len(text), match.end() + 30)
                context = text[context_start:context_end].lower()
                
                # Determine date type from context
                if any(word in context for word in ['expir', 'valid until', 'valid till']):
                    date_type = "expiry"
                elif any(word in context for word in ['issue', 'start', 'effective']):
                    date_type = "issue"
                else:
                    date_type = "unknown"
                
                time_calc = calculate_time_difference(parsed_date, current_date)
                
                all_dates.append({
                    "date_found": parsed_date.strftime("%Y-%m-%d"),
                    "original_text": match.group(1),
                    "date_type": date_type,
                    "relates_to": document_type,
                    "days_from_today": time_calc["total_days"],
                    "status": time_calc["status"],
                    "urgency_level": time_calc["urgency"],
                    "time_calculation": time_calc
                })
            except Exception:
                continue
    
    # Determine overall validity
    current_status = "UNKNOWN"
    expires_on = None
    estimated_expiry = None
    
    # Find expiry dates
    expiry_dates = [d for d in all_dates if d["date_type"] == "expiry"]
    if expiry_dates:
        # Use the earliest expiry date
        expiry_dates.sort(key=lambda x: x["days_from_today"])
        earliest_expiry = expiry_dates[0]
        current_status = earliest_expiry["status"]
        expires_on = earliest_expiry["date_found"]
    else:
        # Try to estimate from issue dates
        issue_dates = [d for d in all_dates if d["date_type"] == "issue"]
        if issue_dates:
            # Use the latest issue date
            issue_dates.sort(key=lambda x: x["days_from_today"], reverse=True)
            latest_issue = issue_dates[0]
            
            # Estimate expiry
            issue_date = datetime.strptime(latest_issue["date_found"], "%Y-%m-%d")
            estimated_expiry_date = issue_date + relativedelta(months=typical_validity_months)
            estimated_expiry = estimated_expiry_date.strftime("%Y-%m-%d")
            
            time_calc = calculate_time_difference(estimated_expiry_date, current_date)
            current_status = time_calc["status"]
    
    # Generate recommendations
    recommendations = []
    critical_issues = []
    
    if current_status == "EXPIRED":
        recommendations.append(f"Document has expired - immediate renewal required")
        critical_issues.append(f"{document_type} is expired and not valid")
    elif current_status in ["EXPIRING_VERY_SOON", "EXPIRING_SOON"]:
        recommendations.append(f"Document expiring soon - start renewal process immediately")
        recommendations.append("Contact relevant authority for renewal procedures")
    elif current_status == "UNKNOWN":
        recommendations.append("Unable to determine validity - manual verification required")
        recommendations.append("Check with issuing authority for current status")
    
    return {
        "document_analysis": {
            "document_type": document_type,
            "confidence_score": "medium" if document_type != "unknown" else "low",
            "current_status": current_status,
            "status_reason": f"Determined from {len(all_dates)} date(s) found and typical {document_type} validity period"
        },
        "dates_extracted": all_dates,
        "validity_assessment": {
            "is_currently_valid": current_status == "VALID",
            "days_remaining": None,  # Would need specific expiry date
            "expires_on": expires_on,
            "estimated_expiry": estimated_expiry,
            "confidence_in_assessment": "medium"
        },
        "recommendations": recommendations,
        "critical_issues": critical_issues
    }

# Enhanced semantic search with intelligent date validation
def semantic_search_with_intelligent_validation(
    query: str,
    collection_name: str = "legal_docs",
    top_k: int = 5,
    embedding_model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
    persist_directory: str = VECTORSTORE_PERSIST_DIR,
    model_name: str = None,
) -> str:
    """
    Enhanced semantic search with intelligent date validation and time calculation
    """
    current_date = datetime.now()
    current_date_str = current_date.strftime("%Y-%m-%d")
    
    # Check if query is about validity, expiry, or time-sensitive matters
    query_lower = query.lower()
    is_time_sensitive = any(word in query_lower for word in [
        'valid', 'expired', 'expiry', 'current', 'active', 'still good', 'up to date',
        'in force', 'effective', 'applicable', 'usable', 'renewal', 'renew'
    ])
    
    # Standard semantic search
    embedder = _get_embedding_model(embedding_model_name)
    vectordb = Chroma(persist_directory=persist_directory, embedding_function=embedder, collection_name=collection_name)
    docs = vectordb.similarity_search(query, k=top_k)
    
    context_text = "\n\n".join([f"--- chunk {i} ---\n{d.page_content.strip()[:2000]}" for i, d in enumerate(docs)])
    
    # Perform intelligent date validation if time-sensitive
    validation_results = ""
    if is_time_sensitive and docs:
        try:
            full_text = "\n".join([doc.page_content for doc in docs])
            date_analysis = intelligent_date_extraction_and_validation(full_text)
            
            doc_analysis = date_analysis.get("document_analysis", {})
            validity = date_analysis.get("validity_assessment", {})
            
            validation_results = f"""
INTELLIGENT VALIDITY ANALYSIS (Current Date: {current_date_str}):
• Document Type: {doc_analysis.get('document_type', 'Unknown')}
• Current Status: {doc_analysis.get('current_status', 'Unknown')}
• Is Currently Valid: {'YES' if validity.get('is_currently_valid') else 'NO'}
• Expires On: {validity.get('expires_on', 'Not determined')}
• Estimated Expiry: {validity.get('estimated_expiry', 'Not estimated')}
• Critical Issues: {len(date_analysis.get('critical_issues', []))} found
• Recommendations: {len(date_analysis.get('recommendations', []))} provided
"""
        except Exception as e:
            validation_results = f"\nVALIDITY ANALYSIS: Error occurred ({str(e)}) - manual review recommended"
    
    # Enhanced prompt with time intelligence
    llm = _get_gemini_llm(model_name=model_name, temperature=0.0)
    
    prompt_template = """You are an intelligent legal document assistant with advanced date calculation capabilities.

TODAY'S DATE: {current_date}
CURRENT DAY: {day_name}

DOCUMENT CONTEXT:
{context}

{validation_results}

USER QUESTION: {question}

INSTRUCTIONS:
- Answer based on retrieved content AND intelligent date validation results
- For validity questions, give DEFINITIVE answers: "YES, valid" or "NO, expired" or "UNCLEAR, needs verification"
- Calculate exact time remaining/elapsed when dates are available
- Use typical validity periods for document types when specific dates aren't clear
- If document is expired, clearly state it is NOT VALID and explain consequences
- If expiring soon, provide specific timeline and urgency level
- Give actionable next steps with specific timelines
- Be definitive rather than vague - users need clear answers

RESPONSE FORMAT:
1. Direct answer to the question
2. Supporting evidence from document
3. Time calculation details if relevant
4. Specific action items with deadlines
5. Risk level assessment

Answer:"""

    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question", "current_date", "day_name", "validation_results"])
    chain = LLMChain(llm=llm, prompt=prompt)
    
    response = chain.run({
        "context": context_text,
        "question": query,
        "current_date": current_date_str,
        "day_name": current_date.strftime("%A"),
        "validation_results": validation_results
    })
    
    return response

def _get_embedding_model(embedding_model_name: str = "sentence-transformers/all-MiniLM-L6-v2"):
    """Return a HuggingFaceEmbeddings instance."""
    return HuggingFaceEmbeddings(model_name=embedding_model_name)

def _get_gemini_llm(model_name: str = None, temperature: float = 0.0):
    """Return a ChatGoogleGenerativeAI LLM object."""
    if model_name is None:
        model_name = GEMINI_MODEL
    
    llm = ChatGoogleGenerativeAI(
        model=model_name,
        google_api_key=GEMINI_API_KEY,
        temperature=temperature,
    )
    return llm

def compare_documents(text1: str, text2: str) -> str:
    """
    Produce a simple redline-like markdown showing differences.
    Uses difflib to produce a unified diff and then formats as markdown.
    """
    a_lines = text1.splitlines()
    b_lines = text2.splitlines()
    diff = difflib.unified_diff(a_lines, b_lines, fromfile="Document A", tofile="Document B", lineterm="")
    md_lines = ["### Document Comparison (unified diff)\n", "```diff"]
    md_lines.extend(list(diff))
    md_lines.append("```")
    return "\n".join(md_lines)


# Additional helper: if you want to embed and return raw embeddings (not used by default)
def embed_texts(texts: List[str], embedding_model_name: str = "sentence-transformers/all-MiniLM-L6-v2"):
    embedder = _get_embedding_model(embedding_model_name)
    vectors = embedder.embed_documents(texts)
    return vectors


# If run as script, example usage (for local dev)
if __name__ == "__main__":
    # Dummy run to ensure imports work
    print("Backend module loaded. VECTORSTORE persist dir:", VECTORSTORE_PERSIST_DIR)